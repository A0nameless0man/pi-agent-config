#!/usr/bin/env python3
"""
Generate test DOCX files demonstrating TOC styles and multi-level heading numbering.

Outputs 3 DOCX files to the Desktop:
  1. legal_numbering_sample.docx      — Standard legal numbering (1., 1.1, 1.1.1) with TOC
  2. chinese_chapter_numbering_sample.docx — Chinese chapter numbering (第一章, 1.1, 1.1.1) with TOC
  3. mixed_numbering_sample.docx      — Mixed: legal + Chinese chapter + Roman numbering with TOC

Each file includes a cover page, table of contents, multi-level heading numbering,
body text, and a verification notes page at the end.
"""

import os
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

# ── Output path ───────────────────────────────────────────────────────────────
DESKTOP = os.path.join(os.environ.get("USERPROFILE", os.path.expanduser("~")), "Desktop")

# ── Verification notes text (reused across all samples) ───────────────────────
VERIFICATION_TEXT = """\
=== 验证说明 / Verification Notes ===

请打开此文档后检查以下项目：

1. 目录 (TOC)
   - Word 应提示"是否更新域"——点击"是"更新整个目录
   - 目录条目应有正确的缩进层级（一级无缩进，二级缩进，三级更深）
   - 目录条目应有正确的字体（一级加粗大字，二级常规，三级小字灰色）
   - 页码应有前导点（......N）

2. 标题自动编号
   - 所有标题应自动编号，无需手动输入数字
   - 编号层级正确
   - 同级编号连续递增
   - 高级别递增时，低级别自动重新从1开始

3. 页面排版
   - 页边距合适，内容不溢出
   - 中文字体显示正常（无方框）
   - 段落间距均匀

如有任何不符合预期的表现，请截图并记录。"""


# ═══════════════════════════════════════════════════════════════════════════════
#  OOXML HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def create_numbering(doc, level_configs):
    """Create a multi-level numbering definition and return its numId.

    Args:
        doc: python-docx Document object.
        level_configs: list of dicts, each with:
            - ilvl (int): level index (0-based)
            - numFmt (str): OOXML number format ('decimal', 'chineseCounting', etc.)
            - lvlText (str): level text template ('%1.', '%1.%2', '第%1章' etc.)
            - start (int, optional): starting number, defaults to 1
            - isLgl (bool, optional): force Arabic numerals for cross-references

    Returns:
        int: numId to use with apply_numbering().
    """
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Get next abstractNum ID
    existing_abs = numbering_elm.xpath("./w:abstractNum/@w:abstractNumId")
    next_abs_id = max([int(x) for x in existing_abs], default=-1) + 1

    # Create abstractNum element
    abstractNum = OxmlElement("w:abstractNum")
    abstractNum.set(qn("w:abstractNumId"), str(next_abs_id))

    # Mark as multi-level
    multiLevelType = OxmlElement("w:multiLevelType")
    multiLevelType.set(qn("w:val"), "multilevel")
    abstractNum.append(multiLevelType)

    for cfg in level_configs:
        ilvl = cfg["ilvl"]
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))

        # Start at
        start_el = OxmlElement("w:start")
        start_el.set(qn("w:val"), str(cfg.get("start", 1)))
        lvl.append(start_el)

        # Number format
        numFmt_el = OxmlElement("w:numFmt")
        numFmt_el.set(qn("w:val"), cfg["numFmt"])
        lvl.append(numFmt_el)

        # Level text template
        lvlText_el = OxmlElement("w:lvlText")
        lvlText_el.set(qn("w:val"), cfg["lvlText"])
        lvl.append(lvlText_el)

        # Legal numbering (forces Arabic for cross-refs)
        if cfg.get("isLgl"):
            isLgl_el = OxmlElement("w:isLgl")
            lvl.append(isLgl_el)

        # Left alignment
        lvlJc_el = OxmlElement("w:lvlJc")
        lvlJc_el.set(qn("w:val"), "left")
        lvl.append(lvlJc_el)

        # Indentation
        pPr_el = OxmlElement("w:pPr")
        ind_el = OxmlElement("w:ind")
        ind_el.set(qn("w:left"), str(420 * (ilvl + 1)))
        ind_el.set(qn("w:hanging"), "420")
        pPr_el.append(ind_el)
        lvl.append(pPr_el)

        abstractNum.append(lvl)

    # Insert abstractNum at the beginning (before num elements)
    numbering_elm.insert(0, abstractNum)

    # Create num instance referencing the abstractNum
    existing_nums = numbering_elm.xpath("./w:num/@w:numId")
    next_num_id = max([int(x) for x in existing_nums], default=0) + 1

    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), str(next_num_id))
    abstractNumId_el = OxmlElement("w:abstractNumId")
    abstractNumId_el.set(qn("w:val"), str(next_abs_id))
    num_el.append(abstractNumId_el)
    numbering_elm.append(num_el)

    return next_num_id


def apply_numbering(paragraph, num_id, ilvl):
    """Apply multi-level heading numbering to a paragraph.

    Args:
        paragraph: docx.text.paragraph.Paragraph object.
        num_id: int from create_numbering().
        ilvl: level index (0 = Heading 1, 1 = Heading 2, etc.).
    """
    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing numPr if any
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(numId_el)
    pPr.insert(0, numPr)


def apply_numbering_to_style(doc, style_name, num_id, ilvl):
    """Inject numbering into a paragraph style definition (styles.xml).

    Once applied, ALL paragraphs using this style are auto-numbered —
    no per-paragraph apply_numbering() calls needed.

    Args:
        doc: Document object.
        style_name: 'Heading 1', 'Heading 2', 'Heading 3', etc.
        num_id: Numbering instance ID from create_numbering().
        ilvl: Level index (0=Heading 1, 1=Heading 2, etc.).
    """
    style = doc.styles[style_name]
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style.element.insert(0, pPr)

    # Remove existing numPr
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(numId_el)
    pPr.insert(0, numPr)


def insert_toc(paragraph):
    """Insert a TOC field into a paragraph using OOXML field codes.

    The TOC will prompt the user to update when the document is opened in Word.
    """
    # Begin field
    run_begin = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    run_begin.append(fc1)
    paragraph._p.append(run_begin)

    # Field instruction
    run_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr.append(instr)
    paragraph._p.append(run_instr)

    # Separator
    run_sep = OxmlElement("w:r")
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    run_sep.append(fc2)
    paragraph._p.append(run_sep)

    # Placeholder text (shown until TOC is updated)
    run_ph = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = '请在Word中右键点击此处，选择"更新域"生成目录'
    run_ph.append(t)
    paragraph._p.append(run_ph)

    # End field
    run_end = OxmlElement("w:r")
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    run_end.append(fc3)
    paragraph._p.append(run_end)


def enable_auto_update(doc):
    """Set w:updateFields=true so Word prompts to update TOC on document open."""
    settings_el = doc.settings._element
    existing = settings_el.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
    else:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings_el.append(uf)


def configure_toc_styles(doc):
    """Configure TOC 1, TOC 2, TOC 3 styles with font, indent, and dot leaders."""
    configs = {
        "TOC 1": {
            "font_name": "SimHei",
            "size": Pt(14),
            "bold": True,
            "color": RGBColor(0, 51, 102),
            "indent": Inches(0),
        },
        "TOC 2": {
            "font_name": "Microsoft YaHei",
            "size": Pt(12),
            "bold": False,
            "color": RGBColor(51, 51, 51),
            "indent": Inches(0.4),
        },
        "TOC 3": {
            "font_name": "Microsoft YaHei",
            "size": Pt(10.5),
            "bold": False,
            "color": RGBColor(102, 102, 102),
            "indent": Inches(0.8),
        },
    }
    for name, cfg in configs.items():
        # Get or create style
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

        # Font
        style.font.name = cfg["font_name"]
        style.font.size = cfg["size"]
        style.font.bold = cfg["bold"]
        style.font.color.rgb = cfg["color"]

        # Set CJK font via OOXML
        rPr = style.element.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                rFonts.set(qn("w:eastAsia"), cfg["font_name"])

        # Indent
        style.paragraph_format.left_indent = cfg["indent"]

        # Dot leader tab stop at 6.0 inches
        tab_stops = style.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_cjk_font(run, font_name="Microsoft YaHei"):
    """Set East Asian font for a run to prevent CJK character boxes."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def apply_global_spacing(doc):
    """Apply 1.3x line spacing to Normal and Heading styles."""
    for i in range(1, 5):
        try:
            heading_style = doc.styles[f"Heading {i}"]
            heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            heading_style.paragraph_format.line_spacing = 1.3
        except KeyError:
            pass
    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    # Set default CJK font for Normal style
    normal.font.name = "Microsoft YaHei"
    rPr = normal.element.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # Configure heading styles for CJK
    heading_fonts = {
        1: ("SimHei", Pt(22), True),
        2: ("Microsoft YaHei", Pt(16), True),
        3: ("Microsoft YaHei", Pt(14), True),
    }
    for level, (font, size, bold) in heading_fonts.items():
        try:
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = font
            hs.font.size = size
            hs.font.bold = bold
            hs.font.color.rgb = RGBColor(0, 51, 102)
            h_rPr = hs.element.find(qn("w:rPr"))
            if h_rPr is not None:
                h_rFonts = h_rPr.find(qn("w:rFonts"))
                if h_rFonts is None:
                    h_rFonts = OxmlElement("w:rFonts")
                    h_rPr.insert(0, h_rFonts)
                h_rFonts.set(qn("w:eastAsia"), font)
        except KeyError:
            pass


def add_cover_page(doc, title_text):
    """Add a centered cover page with a title."""
    # Vertical spacing to center
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    set_cjk_font(run, "SimHei")

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("TOC & Multi-Level Numbering Verification Sample")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(102, 102, 102)
    set_cjk_font(run2)

    doc.add_page_break()


def add_toc_page(doc):
    """Add a TOC page with heading and TOC field."""
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run("目  录")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    set_cjk_font(run, "SimHei")

    doc.add_paragraph("")  # spacing
    toc_para = doc.add_paragraph()
    insert_toc(toc_para)

    doc.add_page_break()


def add_heading_with_body(doc, text, level):
    """Add a heading with 2-3 sentences of body text below it.

    Args:
        doc: Document object.
        text: Heading text (without number prefix).
        level: Heading level (1-3).

    Returns:
        Paragraph: the heading paragraph.
    """
    heading = doc.add_heading(text, level=level)

    # Body text
    body_texts = {
        1: [
            "本章节提供了该主题的总体概述与核心要点，为后续深入讨论奠定基础。",
            "相关内容涵盖了理论框架、实践应用及行业最佳实践等多个维度。",
            "通过系统化的分析，帮助读者建立对该领域的全面认识。",
        ],
        2: [
            "本小节进一步细化主题内容，从具体角度展开详细论述。",
            "结合实际案例和数据，深入剖析关键环节的运作机制与核心逻辑。",
        ],
        3: [
            "本条目聚焦于具体的技术细节或操作要点，提供精确的指导说明。",
            "相关规范和标准在此予以明确，作为实施过程中的参考依据。",
        ],
    }
    sentences = body_texts.get(level, body_texts[1])
    for sentence in sentences:
        p = doc.add_paragraph(sentence)
        for run_obj in p.runs:
            set_cjk_font(run_obj)
    return heading


def add_verification_page(doc):
    """Add a verification notes page at the end of the document."""
    doc.add_page_break()
    for line in VERIFICATION_TEXT.strip().split("\n"):
        p = doc.add_paragraph(line)
        for run in p.runs:
            set_cjk_font(run)
            if line.startswith("==="):
                run.bold = True
                run.font.size = Pt(14)


# ═══════════════════════════════════════════════════════════════════════════════
#  SAMPLE 1: Legal Numbering (1., 1.1, 1.1.1)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_sample_1():
    """Generate legal_numbering_sample.docx on the Desktop."""
    doc = Document()

    # Global styling
    apply_global_spacing(doc)
    configure_toc_styles(doc)
    enable_auto_update(doc)

    # Create legal numbering: 1., 1.1, 1.1.1
    legal_num_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "decimal", "lvlText": "%1.", "start": 1},
        {"ilvl": 1, "numFmt": "decimal", "lvlText": "%1.%2", "start": 1},
        {"ilvl": 2, "numFmt": "decimal", "lvlText": "%1.%2.%3", "start": 1},
    ])

    # Inject numbering into heading styles — once, globally
    apply_numbering_to_style(doc, "Heading 1", legal_num_id, 0)
    apply_numbering_to_style(doc, "Heading 2", legal_num_id, 1)
    apply_numbering_to_style(doc, "Heading 3", legal_num_id, 2)

    # Cover page
    add_cover_page(doc, "法律编号样式测试")

    # TOC page
    add_toc_page(doc)

    # ── Content ──
    # 1. 项目概述
    add_heading_with_body(doc, "项目概述", 1)

    #   1.1 项目背景
    add_heading_with_body(doc, "项目背景", 2)
    #     1.1.1 市场环境
    add_heading_with_body(doc, "市场环境", 3)
    #     1.1.2 政策法规
    add_heading_with_body(doc, "政策法规", 3)

    #   1.2 项目目标
    add_heading_with_body(doc, "项目目标", 2)

    # 2. 技术方案
    add_heading_with_body(doc, "技术方案", 1)

    #   2.1 总体架构
    add_heading_with_body(doc, "总体架构", 2)
    #   2.2 关键技术
    add_heading_with_body(doc, "关键技术", 2)
    #     2.2.1 核心技术A
    add_heading_with_body(doc, "核心技术A", 3)
    #     2.2.2 核心技术B
    add_heading_with_body(doc, "核心技术B", 3)

    # 3. 实施计划
    add_heading_with_body(doc, "实施计划", 1)

    #   3.1 阶段划分
    add_heading_with_body(doc, "阶段划分", 2)
    #   3.2 资源配置
    add_heading_with_body(doc, "资源配置", 2)

    # Verification page
    add_verification_page(doc)

    # Save
    output_path = os.path.join(DESKTOP, "legal_numbering_sample.docx")
    doc.save(output_path)
    print(f"[OK] {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
#  SAMPLE 2: Chinese Chapter Numbering (第一章, 1.1, 1.1.1)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_sample_2():
    """Generate chinese_chapter_numbering_sample.docx on the Desktop."""
    doc = Document()

    # Global styling
    apply_global_spacing(doc)
    configure_toc_styles(doc)
    enable_auto_update(doc)

    # Create Chinese chapter numbering: 第X章, X.1, X.1.1
    # Sub-levels use isLgl=True so cross-references (%1) render as decimal,
    # producing "1.1" instead of "一.1".
    ch_num_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "chineseCounting", "lvlText": "第%1章", "start": 1},
        {"ilvl": 1, "numFmt": "decimal", "lvlText": "%1.%2", "start": 1, "isLgl": True},
        {"ilvl": 2, "numFmt": "decimal", "lvlText": "%1.%2.%3", "start": 1, "isLgl": True},
    ])

    # Inject numbering into heading styles
    apply_numbering_to_style(doc, "Heading 1", ch_num_id, 0)
    apply_numbering_to_style(doc, "Heading 2", ch_num_id, 1)
    apply_numbering_to_style(doc, "Heading 3", ch_num_id, 2)

    # Cover page
    add_cover_page(doc, "中文章节编号样式测试")

    # TOC page
    add_toc_page(doc)

    # ── Content ──
    # 第一章 项目背景
    add_heading_with_body(doc, "项目背景", 1)
    #   1.1 行业现状
    add_heading_with_body(doc, "行业现状", 2)
    #     1.1.1 国内市场
    add_heading_with_body(doc, "国内市场", 3)
    #     1.1.2 国际市场
    add_heading_with_body(doc, "国际市场", 3)
    #   1.2 发展趋势
    add_heading_with_body(doc, "发展趋势", 2)

    # 第二章 技术研究
    add_heading_with_body(doc, "技术研究", 1)
    #   2.1 系统设计
    add_heading_with_body(doc, "系统设计", 2)
    #     2.1.1 硬件架构
    add_heading_with_body(doc, "硬件架构", 3)
    #     2.1.2 软件架构
    add_heading_with_body(doc, "软件架构", 3)
    #   2.2 接口规范
    add_heading_with_body(doc, "接口规范", 2)

    # 第三章 实施与运维
    add_heading_with_body(doc, "实施与运维", 1)
    #   3.1 部署方案
    add_heading_with_body(doc, "部署方案", 2)
    #   3.2 监控体系
    add_heading_with_body(doc, "监控体系", 2)

    # Verification page
    add_verification_page(doc)

    # Save
    output_path = os.path.join(DESKTOP, "chinese_chapter_numbering_sample.docx")
    doc.save(output_path)
    print(f"[OK] {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
#  SAMPLE 3: Mixed Numbering (Legal + Chinese Chapter + Roman + Body Lists)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_sample_3():
    """Generate mixed_numbering_sample.docx on the Desktop.

    Contains three sections with different numbering schemes:
      - Section A: Legal numbering (1., 1.1, 1.1.1)
      - Section B: Chinese chapter numbering (第X章, X.1, X.1.1)
      - Section C: Roman numeral numbering (I., A., 1.)
    Plus body paragraphs with numbered lists (1), 2), 3)...).
    """
    doc = Document()

    # Global styling
    apply_global_spacing(doc)
    configure_toc_styles(doc)
    enable_auto_update(doc)

    # ── Create numbering definitions ──
    legal_num_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "decimal", "lvlText": "%1.", "start": 1},
        {"ilvl": 1, "numFmt": "decimal", "lvlText": "%1.%2", "start": 1},
        {"ilvl": 2, "numFmt": "decimal", "lvlText": "%1.%2.%3", "start": 1},
    ])

    ch_num_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "chineseCounting", "lvlText": "第%1章", "start": 1},
        {"ilvl": 1, "numFmt": "decimal", "lvlText": "%1.%2", "start": 1, "isLgl": True},
        {"ilvl": 2, "numFmt": "decimal", "lvlText": "%1.%2.%3", "start": 1, "isLgl": True},
    ])

    roman_num_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "upperRoman", "lvlText": "%1.", "start": 1},
        {"ilvl": 1, "numFmt": "upperLetter", "lvlText": "%2.", "start": 1},
        {"ilvl": 2, "numFmt": "decimal", "lvlText": "%3.", "start": 1},
    ])

    # Each list gets its own numbering instance so they restart at 1
    body_list_a_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "decimal", "lvlText": "%1)", "start": 1},
    ])
    body_list_b_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "decimal", "lvlText": "%1)", "start": 1},
    ])
    body_list_c_id = create_numbering(doc, [
        {"ilvl": 0, "numFmt": "decimal", "lvlText": "%1)", "start": 1},
    ])

    # Cover page
    add_cover_page(doc, "混合编号样式测试")

    # TOC page
    add_toc_page(doc)

    # ──────────────────────────────────────────────────────────────────────
    #  SECTION A: Legal Numbering
    # ──────────────────────────────────────────────────────────────────────
    # Section divider (bold paragraph, not heading — avoids picking up numbering)
    section_a = doc.add_paragraph()
    run_a = section_a.add_run("法律编号章节")
    run_a.bold = True
    run_a.font.size = Pt(22)
    run_a.font.color.rgb = RGBColor(0, 51, 102)
    set_cjk_font(run_a, "SimHei")

    # Apply legal numbering to heading styles for this section
    apply_numbering_to_style(doc, "Heading 1", legal_num_id, 0)
    apply_numbering_to_style(doc, "Heading 2", legal_num_id, 1)
    apply_numbering_to_style(doc, "Heading 3", legal_num_id, 2)

    add_heading_with_body(doc, "项目规划", 1)
    add_heading_with_body(doc, "需求分析", 2)
    add_heading_with_body(doc, "可行性评估", 2)
    add_heading_with_body(doc, "执行方案", 1)

    # Numbered list in body text under Section A
    doc.add_paragraph("以下为本节的行动要点：")
    body_list_items_a = [
        "完成项目立项审批流程",
        "组建跨部门协作团队",
        "制定详细实施时间表",
    ]
    for item in body_list_items_a:
        p = doc.add_paragraph(item)
        apply_numbering(p, body_list_a_id, 0)
        for run in p.runs:
            set_cjk_font(run)

    # ──────────────────────────────────────────────────────────────────────
    #  SECTION B: Chinese Chapter Numbering
    # ──────────────────────────────────────────────────────────────────────
    # Section divider
    section_b = doc.add_paragraph()
    run_b = section_b.add_run("中文章节编号章节")
    run_b.bold = True
    run_b.font.size = Pt(22)
    run_b.font.color.rgb = RGBColor(0, 51, 102)
    set_cjk_font(run_b, "SimHei")

    # Apply Chinese chapter numbering to heading styles for this section
    apply_numbering_to_style(doc, "Heading 1", ch_num_id, 0)
    apply_numbering_to_style(doc, "Heading 2", ch_num_id, 1)
    apply_numbering_to_style(doc, "Heading 3", ch_num_id, 2)

    add_heading_with_body(doc, "战略规划", 1)
    add_heading_with_body(doc, "目标设定", 2)
    add_heading_with_body(doc, "短期目标", 3)
    add_heading_with_body(doc, "长期愿景", 3)
    add_heading_with_body(doc, "路径选择", 2)

    # Numbered list under Section B
    doc.add_paragraph("关键里程碑：")
    body_list_items_b = [
        "第一季度完成初步调研",
        "第二季度启动试点项目",
        "第三季度全面推广实施",
        "第四季度总结评估优化",
    ]
    for item in body_list_items_b:
        p = doc.add_paragraph(item)
        apply_numbering(p, body_list_b_id, 0)
        for run in p.runs:
            set_cjk_font(run)

    # ──────────────────────────────────────────────────────────────────────
    #  SECTION C: Roman Numeral Numbering
    # ──────────────────────────────────────────────────────────────────────
    # Section divider
    section_c = doc.add_paragraph()
    run_c = section_c.add_run("罗马数字编号章节")
    run_c.bold = True
    run_c.font.size = Pt(22)
    run_c.font.color.rgb = RGBColor(0, 51, 102)
    set_cjk_font(run_c, "SimHei")

    # Apply Roman numbering to heading styles for this section
    apply_numbering_to_style(doc, "Heading 1", roman_num_id, 0)
    apply_numbering_to_style(doc, "Heading 2", roman_num_id, 1)
    apply_numbering_to_style(doc, "Heading 3", roman_num_id, 2)

    add_heading_with_body(doc, "国际标准参考", 1)
    add_heading_with_body(doc, "标准概述", 2)
    add_heading_with_body(doc, "实施要点", 3)
    add_heading_with_body(doc, "验收标准", 3)
    add_heading_with_body(doc, "合规要求", 2)

    # Numbered list under Section C
    doc.add_paragraph("检查清单：")
    body_list_items_c = [
        "确认所有文件已签署",
        "验证系统配置完整性",
        "执行最终验收测试",
    ]
    for item in body_list_items_c:
        p = doc.add_paragraph(item)
        apply_numbering(p, body_list_c_id, 0)
        for run in p.runs:
            set_cjk_font(run)

    # Verification page
    add_verification_page(doc)

    # Save
    output_path = os.path.join(DESKTOP, "mixed_numbering_sample.docx")
    doc.save(output_path)
    print(f"[OK] {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("Generating test DOCX files...")
    print()

    try:
        path1 = generate_sample_1()
        path2 = generate_sample_2()
        path3 = generate_sample_3()
    except Exception as e:
        print(f"\n  [ERROR] {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)

    print()
    print("Done! 3 DOCX files written to Desktop:")
    print(f"  1. {path1}")
    print(f"  2. {path2}")
    print(f"  3. {path3}")
    print()
    print("Open each file in Microsoft Word and:")
    print("  - Click 'Yes' when prompted to update fields (this generates the TOC)")
    print("  - Verify heading numbering matches expectations")
    print("  - Check the verification notes page at the end of each document")


if __name__ == "__main__":
    main()
